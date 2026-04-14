"""TZ generation module with source-faithful drafting and optional AI rewrite."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Optional

from docx import Document

from src.nlp.ai_client import AIClient

SECTION_SPECS = [
    ("1.1", "Наименование приоритета программы", ["приоритет", "направлен", "область"]),
    ("1.2", "Наименование специализированного направления", ["специализирован", "технолог", "тематика"]),
    ("2.1", "Цель программы", ["цель", "целью", "назначение"]),
    ("2.2", "Задачи программы", ["задач", "этап", "необходимо", "требуется"]),
    ("3", "Пункты стратегических и программных документов", ["документ", "стратег", "программа", "постановл"]),
    ("4.1", "Прямые результаты", ["результат", "публикац", "патент", "прототип"]),
    ("4.2", "Конечный результат", ["эффект", "внедрен", "итог", "конечн"]),
    ("5", "Предельная сумма программы", ["бюджет", "тенге", "сумм", "финансир"]),
]


def generate_tz(payload: dict[str, Any]) -> dict[str, Any]:
    context = payload.get("context") or {}
    source_text = _normalize_whitespace(payload.get("text") or "")
    title_hint = str(payload.get("project_title") or "").strip()

    base = _build_extractive_draft(source_text=source_text, title_hint=title_hint, payload=payload)
    generated_text = base["generated_text"]
    method = "extractive-faithful"
    ai_meta: dict[str, Any] = {"status": "not_used", "provider": None}

    ai_provider = str(context.get("ai_provider", "")).strip() or None
    ai_api_key = str(context.get("ai_api_key", "")).strip() or None
    if source_text:
        ai_result = _try_ai_rewrite(
            source_text=source_text,
            draft_text=generated_text,
            ai_provider=ai_provider,
            ai_api_key=ai_api_key,
        )
        if ai_result.get("status") == "ready":
            generated_text = str(ai_result.get("generated_text", "")).strip() or generated_text
            method = "ai-assisted-faithful"
        ai_meta = ai_result

    return {
        "agent": "generation_agent",
        "generated_title": base["generated_title"],
        "template": "pcf_scientific_tz",
        "generated_text": generated_text,
        "generation_method": method,
        "source_text_length": len(source_text),
        "section_coverage": base["section_coverage"],
        "ai_generation": ai_meta,
    }


def render_generated_tz_docx(title: str, generated_text: str, output_path: Path) -> Path:
    document = Document()
    document.add_heading(title or "Сформированное ТЗ", level=0)

    for raw_line in (generated_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+(\.\d+)?\.", line):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.bold = True
        else:
            document.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _build_extractive_draft(source_text: str, title_hint: str, payload: dict[str, Any]) -> dict[str, Any]:
    if not source_text:
        fallback = _fallback_draft(payload)
        fallback["section_coverage"] = {section_id: "missing" for section_id, *_ in SECTION_SPECS}
        fallback["source_quality_note"] = "empty_source_text"
        return fallback

    sentences = _split_sentences(source_text)
    generated_title = _guess_title(source_text, title_hint) or "Проект технического задания"
    lines = [
        "1. Общие сведения",
    ]
    coverage: dict[str, str] = {}

    for section_id, section_title, keywords in SECTION_SPECS:
        section_header = f"{section_id}. {section_title}"
        lines.append(section_header)
        extracted = _extract_relevant_sentences(sentences=sentences, keywords=keywords, limit=3)
        if extracted:
            coverage[section_id] = "present"
            for item in extracted:
                if section_id == "2.2":
                    lines.append(f"- {item}")
                else:
                    lines.append(item)
        else:
            coverage[section_id] = "missing"
            lines.append("Требуется уточнение по исходному документу.")
        lines.append("")

    return {
        "generated_title": generated_title,
        "generated_text": "\n".join(lines).strip(),
        "section_coverage": coverage,
    }


def _fallback_draft(payload: dict[str, Any]) -> dict[str, Any]:
    project_title = payload.get("project_title", "Новая научно-техническая программа")
    priority = payload.get("priority", "Энергия, передовые материалы и транспорт")
    specialization = payload.get("specialization", "Цифровые технологии и ИИ")
    budget_total = payload.get("budget_total", "1 500 000 тыс. тенге")
    years = payload.get("years", ["2026", "2027", "2028"])
    year_budget = payload.get("year_budget", "500 000 тыс. тенге")
    generated_text = f"""
1. Общие сведения
1.1. Наименование приоритета программы: {priority}
1.2. Наименование специализированного направления: {specialization}

2. Цели и задачи программы
2.1. Цель программы: {project_title}
2.2. Задачи программы:
- Уточнить структуру ТЗ на основе исходного документа
- Добавить измеримые KPI и план внедрения

3. Пункты стратегических и программных документов
Требуется уточнение по исходному документу.

4. Ожидаемые результаты
4.1. Прямые результаты: Требуется уточнение по исходному документу.
4.2. Конечный результат: Требуется уточнение по исходному документу.

5. Предельная сумма программы
На весь срок: {budget_total}
{chr(10).join([f"- на {y} г.: {year_budget}" for y in years])}
""".strip()
    return {"generated_title": project_title, "generated_text": generated_text}


def _try_ai_rewrite(
    source_text: str,
    draft_text: str,
    ai_provider: Optional[str],
    ai_api_key: Optional[str],
) -> dict[str, Any]:
    ai_client = AIClient(provider=ai_provider, api_key=ai_api_key)
    if not ai_client.is_available():
        return {
            "status": "unavailable",
            "provider": ai_client.provider,
            "summary": f"{ai_client.provider.upper()} API не настроен. Использован extractive режим.",
        }

    prompt = (
        "Ты редактор ТЗ. Перепиши черновик по шаблону, но строго на основе исходного текста. "
        "Нельзя выдумывать факты, цифры, документы или KPI. Если данных нет, пиши: "
        "\"Требуется уточнение по исходному документу.\".\n\n"
        f"Исходный текст ТЗ:\n{source_text[:8000]}\n\n"
        f"Черновик:\n{draft_text}\n\n"
        "Верни только JSON:\n"
        "{"
        "\"generated_title\":\"...\","
        "\"generated_text\":\"...\","
        "\"fidelity_note\":\"кратко как сохранялась связь с исходным текстом\""
        "}"
    )

    try:
        parsed = _call_ai_json(ai_client, prompt)
    except Exception as exc:
        return {
            "status": "error",
            "provider": ai_client.provider,
            "summary": "AI-переписывание не удалось. Использован extractive режим.",
            "error": str(exc),
        }

    generated_text = str(parsed.get("generated_text", "")).strip()
    if not generated_text:
        return {
            "status": "error",
            "provider": ai_client.provider,
            "summary": "AI вернул пустой текст. Использован extractive режим.",
            "error": "empty_generated_text",
        }

    return {
        "status": "ready",
        "provider": ai_client.provider,
        "model": getattr(ai_client, "model", ""),
        "generated_title": str(parsed.get("generated_title", "")).strip(),
        "generated_text": generated_text,
        "fidelity_note": str(parsed.get("fidelity_note", "")).strip(),
    }


def _call_ai_json(ai_client: AIClient, prompt: str) -> dict[str, Any]:
    provider = ai_client.provider
    if provider in {"openai", "github", "openrouter"}:
        response = ai_client.client.chat.completions.create(
            model=ai_client.model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Ты аккуратный редактор ТЗ. Строго сохраняй факты исходника. "
                        "Отвечай только валидным JSON."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=3000,
            temperature=0.15,
            response_format={"type": "json_object"},
        )
        return _safe_json_parse(response.choices[0].message.content)

    if provider == "claude":
        response = ai_client.client.messages.create(
            model=ai_client.model,
            max_tokens=3000,
            temperature=0.15,
            messages=[{"role": "user", "content": prompt}],
        )
        return _safe_json_parse(response.content[0].text)

    raise RuntimeError(f"Unsupported provider: {provider}")


def _safe_json_parse(raw: str) -> dict[str, Any]:
    text = (raw or "").strip()
    if not text:
        raise ValueError("Empty AI response")
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        pass

    if "```json" in text:
        start = text.find("```json") + 7
        end = text.find("```", start)
        text = text[start:end].strip() if end != -1 else text[start:].strip()
    elif "```" in text:
        start = text.find("```") + 3
        end = text.find("```", start)
        text = text[start:end].strip() if end != -1 else text[start:].strip()
    first = text.find("{")
    last = text.rfind("}")
    if first != -1 and last != -1 and last > first:
        text = text[first : last + 1]

    parsed = json.loads(text)
    if not isinstance(parsed, dict):
        raise ValueError("AI response JSON must be an object")
    return parsed


def _normalize_whitespace(text: str) -> str:
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _split_sentences(text: str) -> list[str]:
    raw = re.split(r"(?<=[.!?;])\s+|\n+", text)
    out: list[str] = []
    for item in raw:
        normalized = item.strip()
        if len(normalized) < 20:
            continue
        out.append(normalized)
    return out


def _extract_relevant_sentences(sentences: list[str], keywords: list[str], limit: int = 3) -> list[str]:
    selected: list[str] = []
    seen: set[str] = set()
    for sentence in sentences:
        lowered = sentence.lower()
        if not any(keyword in lowered for keyword in keywords):
            continue
        key = sentence[:120].lower()
        if key in seen:
            continue
        seen.add(key)
        selected.append(sentence)
        if len(selected) >= limit:
            break
    return selected


def _guess_title(source_text: str, title_hint: str) -> str:
    if title_hint:
        return title_hint
    first_lines = [line.strip() for line in source_text.splitlines() if line.strip()]
    for line in first_lines[:6]:
        if 8 <= len(line) <= 160:
            if re.match(r"^\d+(\.\d+)*[.)]?$", line):
                continue
            return line
    return "Проект технического задания"
