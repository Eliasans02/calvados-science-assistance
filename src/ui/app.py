"""Streamlit Dashboard for Dead Regulations Detector."""

import streamlit as st
import json
import sys
import os
import inspect
from dataclasses import dataclass
from typing import List
from pathlib import Path
from datetime import datetime
import time
from collections import Counter
from difflib import SequenceMatcher
from io import BytesIO
import re
import plotly.express as px
import plotly.graph_objects as go
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add parent directory to path
parent_dir = str(Path(__file__).parent.parent)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)

from nlp.ai_client import AIClient
from analysis.dead_reg_detector import DeadRegulationDetector
from scraper.adilet_scraper import AdiletScraper  # Legacy (kept for compatibility)
import config as cfg
from extraction.text_extractor import extract_text_from_uploaded_file
from metrics import StatsEvent, get_stats_service


def _is_valid_github_token(token: str) -> bool:
    """Validate GitHub token format to avoid false-positive AI activation."""
    if not token:
        return False
    token = token.strip()
    valid_prefixes = ("ghp_", "github_pat_", "gho_", "ghu_", "ghs_")
    return token.startswith(valid_prefixes) and len(token) >= 20


SEVERITY_RU = {
    "High": "Высокая",
    "Medium": "Средняя",
    "Low": "Низкая",
}

PRIORITY_RU = {
    "High": "Высокий",
    "Medium": "Средний",
    "Low": "Низкий",
}

ISSUE_TYPE_RU = {
    "outdated_terms": "Устаревшие термины",
    "contradiction": "Противоречие",
    "duplication": "Дублирование",
    "inapplicability": "Неприменимость",
}


def severity_ru(value: str) -> str:
    return SEVERITY_RU.get(value, value or "Не указана")


def priority_ru(value: str) -> str:
    return PRIORITY_RU.get(value, value or "Не указано")


def issue_type_ru(value: str) -> str:
    return ISSUE_TYPE_RU.get(value, (value or "unknown").replace("_", " ").title())


def _format_ai_error_message(err: str, summary: str = "") -> str:
    text = f"{err or ''} {summary or ''}".strip()
    if "RateLimitReached" in text or "429" in text:
        wait_seconds = None
        m = re.search(r"wait\s+(\d+)\s+seconds", text, flags=re.IGNORECASE)
        if m:
            try:
                wait_seconds = int(m.group(1))
            except ValueError:
                wait_seconds = None
        if wait_seconds is not None:
            hours = wait_seconds // 3600
            minutes = (wait_seconds % 3600) // 60
            return (
                f"Лимит AI исчерпан (HTTP 429). Повторите позже "
                f"(примерно через {hours} ч {minutes} мин). "
                "Сейчас выполнен базовый rule-based анализ."
            )
        return "Лимит AI исчерпан (HTTP 429). Сейчас выполнен базовый rule-based анализ."
    if "tokens_limit_reached" in text or "Request body too large" in text or "413" in text:
        return "AI отклонил слишком большой запрос (413). Текст автоматически разбивается на части. Попробуйте запустить анализ ещё раз."
    if err == "AI_NO_ISSUES_OR_FAILED":
        return "AI не вернул результат. Выполнен базовый rule-based анализ."
    return f"AI ошибка: {err}"

# Page config
st.set_page_config(
    page_title="Dead Regulations Detector",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    :root {
        --bg: #f4f7fb;
        --panel: #ffffff;
        --sidebar: #edf2f8;
        --border: #d7e0ea;
        --text: #0f172a;
        --muted: #475569;
        --accent: #1d4ed8;
        --accent-hover: #1e40af;
    }

    html, body, .stApp, [data-testid="stAppViewContainer"], [data-testid="stHeader"] {
        background: var(--bg) !important;
        color: var(--text) !important;
    }

    [data-testid="stSidebar"] {
        background: var(--sidebar) !important;
        border-right: 1px solid var(--border) !important;
    }

    .main-header {
        font-size: 2.3rem;
        font-weight: 750;
        color: var(--text);
        margin-bottom: 0.45rem;
        letter-spacing: -0.4px;
    }

    .subtitle {
        color: var(--muted);
        margin-bottom: 1.4rem;
        font-size: 1.04rem;
        font-weight: 500;
    }

    .stButton>button {
        background: var(--accent) !important;
        color: #ffffff !important;
        border: 1px solid var(--accent) !important;
        border-radius: 10px !important;
        font-weight: 650 !important;
        padding: 0.62rem 1.3rem !important;
        box-shadow: none !important;
    }

    .stButton>button span,
    .stButton>button p,
    .stButton>button div {
        color: #ffffff !important;
    }

    .stButton>button:hover {
        background: var(--accent-hover) !important;
        color: #ffffff !important;
        border-color: var(--accent-hover) !important;
    }

    [data-testid="stMarkdownContainer"] p,
    [data-testid="stMarkdownContainer"] li,
    [data-testid="stMarkdownContainer"] span,
    [data-testid="stMarkdownContainer"] h1,
    [data-testid="stMarkdownContainer"] h2,
    [data-testid="stMarkdownContainer"] h3,
    [data-testid="stMarkdownContainer"] h4,
    [data-testid="stText"],
    [data-testid="stCaptionContainer"] p,
    [data-testid="stSidebar"] * {
        color: var(--text) !important;
    }

    [data-testid="stCaptionContainer"] p {
        color: var(--muted) !important;
    }

    [data-baseweb="input"] > div,
    [data-baseweb="select"] > div,
    textarea,
    input {
        background: var(--panel) !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
        border-radius: 10px !important;
    }

    textarea[disabled],
    textarea[readonly],
    [data-baseweb="textarea"] textarea {
        color: #0f172a !important;
        -webkit-text-fill-color: #0f172a !important;
        opacity: 1 !important;
        background: #ffffff !important;
    }

    input::placeholder, textarea::placeholder {
        color: #64748b !important;
        opacity: 1 !important;
    }

    [data-baseweb="select"] svg {
        fill: var(--text) !important;
    }

    [data-testid="stTextInput"] button,
    [data-testid="stNumberInput"] button {
        background: #f8fafc !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
    }

    [data-testid="stTextInput"] button:hover,
    [data-testid="stNumberInput"] button:hover {
        background: #eef2f7 !important;
        color: var(--text) !important;
    }

    [data-testid="stFileUploaderDropzone"] {
        background: var(--panel) !important;
        border: 1px dashed #94a3b8 !important;
        border-radius: 12px !important;
    }

    [data-testid="stFileUploaderDropzone"] * {
        color: var(--text) !important;
        fill: var(--text) !important;
    }

    [data-testid="stFileUploaderDropzone"] button {
        background: #f8fafc !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
    }

    .streamlit-expanderHeader,
    .dataframe,
    [data-testid="stDataFrame"] {
        background: var(--panel) !important;
        color: var(--text) !important;
        border-color: var(--border) !important;
    }

    div[data-testid="stMetricValue"] {
        color: var(--text) !important;
    }

    div[data-testid="stMetricLabel"] {
        color: #334155 !important;
        font-weight: 700 !important;
    }
</style>
""", unsafe_allow_html=True)


def main():
    """Main application."""
    init_session_state()
    
    # Header
    st.markdown('<div class="main-header">Dead Regulations Detector</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Система анализа нормативно-правовых актов Республики Казахстан</div>', unsafe_allow_html=True)
    st.markdown("---")
    
    # Sidebar
    with st.sidebar:
        st.header("Параметры анализа")
        
        analysis_mode = st.radio(
            "Режим работы",
            [
                "Загрузить документ",
                "Мульти-анализ (PDF/DOCX/TXT)",
                "Поиск в базе adilet.zan.kz",
                "Демонстрационные примеры",
                "Аналитика системы",
            ]
        )
        
        # AI Provider selection
        ai_provider = st.selectbox(
            "Режим AI анализа",
            ["github", "openrouter", "openai", "claude", "none"],
            index=(
                0 if cfg.AI_PROVIDER == "github"
                else 1 if cfg.AI_PROVIDER == "openrouter"
                else 2 if cfg.AI_PROVIDER == "openai"
                else 3 if cfg.AI_PROVIDER == "claude"
                else 4
            ),
            help="Выберите провайдера AI или отключите AI анализ"
        )
        
        use_ai = ai_provider != "none"
        
        if use_ai:
            api_key = ""
            if ai_provider == "openrouter":
                api_key = st.text_input(
                    "OpenRouter API Key",
                    type="password",
                    value=cfg.OPENROUTER_API_KEY if hasattr(cfg, "OPENROUTER_API_KEY") else "",
                    help="Получить: https://openrouter.ai/keys"
                )
                if not api_key:
                    st.warning("Добавьте OpenRouter API ключ для AI анализа")
                st.caption(f"Модель: {getattr(cfg, 'OPENROUTER_MODEL', 'moonshotai/kimi-k2:free')}")
            elif ai_provider == "openai":
                api_key = st.text_input(
                    "OpenAI API Key", 
                    type="password", 
                    value=cfg.OPENAI_API_KEY if cfg.OPENAI_API_KEY != "your-openai-key-here" else "",
                    help="Получить: https://platform.openai.com/api-keys"
                )
                if not api_key or api_key == "your-openai-key-here":
                    st.warning("Добавьте OpenAI API ключ для AI анализа")
            else:
                if ai_provider == "github":
                    api_key = st.text_input(
                        "GitHub Token",
                        type="password",
                        value=cfg.GITHUB_TOKEN if hasattr(cfg, "GITHUB_TOKEN") else "",
                        help="Получить: gh auth token"
                    )
                    if not api_key:
                        st.warning("Добавьте GitHub token для AI анализа")
                    elif not _is_valid_github_token(api_key):
                        st.error("Неверный формат GitHub token. Используйте токен вида ghp_... или github_pat_...")
                        api_key = ""
                else:
                    api_key = st.text_input(
                        "Claude API Key", 
                        type="password", 
                        value=cfg.CLAUDE_API_KEY if cfg.CLAUDE_API_KEY != "sk-ant-your-key-here" else "",
                        help="Получить: https://console.anthropic.com/"
                    )
                    if not api_key or api_key == "sk-ant-your-key-here":
                        st.warning("Добавьте Claude API ключ для AI анализа")
            st.session_state.runtime_ai_key = api_key
            st.session_state.runtime_ai_provider = ai_provider
            if api_key:
                st.success("API ключ подключен")
            else:
                st.info("AI анализ недоступен")
        else:
            st.session_state.runtime_ai_key = ""
            st.session_state.runtime_ai_provider = "none"
        
        st.markdown("---")
        st.markdown("### Статистика системы")
        stats = get_stats_service().get_stats()
        st.metric("Проанализировано документов", str(stats.get("total_processed_documents", 0)))
        st.metric("Выявлено проблем", str(st.session_state.total_issues))
        st.metric("Успешных анализов", str(stats.get("analysis_results", {}).get("success", 0)))
        st.metric("Ошибок анализа", str(stats.get("analysis_results", {}).get("failed", 0)))
        st.caption(f"Текущий статус: {st.session_state.last_status}")
        with st.expander("История анализов", expanded=False):
            history = list(reversed(stats.get("recent_analyses", [])))
            if not history:
                st.caption("История пока пуста")
            else:
                history_rows = []
                for item in history[:50]:
                    history_rows.append(
                        {
                            "Время": item.get("timestamp", ""),
                            "Документ": item.get("title", "Без названия"),
                            "Тип": item.get("document_type", "unknown"),
                            "Источник": item.get("source", "unknown"),
                            "Статус": item.get("status", "success"),
                            "AI": item.get("ai_result", "skipped"),
                            "Проблем": item.get("issues_found", 0),
                            "Ошибка": item.get("error", ""),
                        }
                    )
                st.dataframe(pd.DataFrame(history_rows), use_container_width=True, hide_index=True)
    
    # Main content
    if analysis_mode == "Загрузить документ":
        show_text_input(use_ai, ai_provider if use_ai else None)
    elif analysis_mode == "Мульти-анализ (PDF/DOCX/TXT)":
        show_multi_analysis(use_ai, ai_provider if use_ai else None)
    elif analysis_mode == "Поиск в базе adilet.zan.kz":
        show_search_interface(use_ai, ai_provider if use_ai else None)
    elif analysis_mode == "Аналитика системы":
        show_system_analytics()
    else:
        show_demo_examples(use_ai, ai_provider if use_ai else None)


def show_text_input(use_ai: bool, ai_provider: str = None):
    """Show text input interface."""
    st.header("Анализ текста документа")
    version_compare = st.checkbox("Сравнить две версии документа")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        doc_title = st.text_input("Название документа", placeholder="Например: Закон о...")
        uploaded_file = st.file_uploader(
            "Файл документа",
            type=["txt", "pdf", "docx"],
            help="Поддерживаются TXT, PDF, DOCX"
        )
        doc_text = st.text_area(
            "Текст документа",
            height=300,
            placeholder="Вставьте текст нормативного акта для анализа..."
        )
        old_version_text = ""
        old_uploaded_file = None
        if version_compare:
            old_uploaded_file = st.file_uploader(
                "Файл предыдущей версии",
                type=["txt", "pdf", "docx"],
                help="Можно сравнить файл с файлом"
            )
            old_version_text = st.text_area(
                "Предыдущая версия документа",
                height=220,
                placeholder="Вставьте старую версию для сравнения..."
            )
    
    with col2:
        st.info("""
        Что ищет система:
        - Устаревшие термины
        - Противоречия
        - Дублирование
        - Неприменимые нормы
        """)
    
    file_text = ""
    old_file_text = ""
    if uploaded_file is not None:
        file_text = extract_text_from_file(uploaded_file)
        if file_text:
            st.caption(f"Загружен файл: {uploaded_file.name} ({len(file_text)} символов)")
    if version_compare and old_uploaded_file is not None:
        old_file_text = extract_text_from_file(old_uploaded_file)
        if old_file_text:
            st.caption(f"Загружен файл прошлой версии: {old_uploaded_file.name} ({len(old_file_text)} символов)")

    final_text = doc_text.strip() if doc_text.strip() else file_text.strip()
    final_old_text = old_version_text.strip() if old_version_text.strip() else old_file_text.strip()
    
    # Show info about text sources
    if final_text and uploaded_file:
        st.info(f"Используется текст из файла: {len(final_text):,} символов")
    elif final_text:
        st.info(f"Используется введённый текст: {len(final_text):,} символов")

    if st.button("Анализировать", type="primary"):
        if not final_text:
            st.error("Введите текст документа")
            return

        document = {
            "id": "uploaded",
            "title": doc_title or (uploaded_file.name if uploaded_file else "Загруженный документ"),
            "full_text": final_text,
            "doc_type": (
                Path(uploaded_file.name).suffix.lstrip(".").lower()
                if uploaded_file and getattr(uploaded_file, "name", "")
                else "text"
            ),
            "source": "upload"
        }
        run_analysis_pipeline(document, use_ai, ai_provider)
        if version_compare and final_old_text:
            detector = DeadRegulationDetector()
            
            # Show what we're comparing
            with st.expander("Информация о сравниваемых текстах", expanded=False):
                col_a, col_b = st.columns(2)
                with col_a:
                    st.write(f"**Старая версия:** {len(final_old_text):,} символов")
                    st.caption(f"Строк: {len(final_old_text.splitlines())}")
                with col_b:
                    st.write(f"**Новая версия:** {len(final_text):,} символов")
                    st.caption(f"Строк: {len(final_text.splitlines())}")
                st.caption("Примечание: Выполняется лёгкая нормализация (стандартизация переносов строк)")
            
            cmp = detector.compare_versions(final_old_text, final_text)
            
            st.subheader("Сравнение версий документа")
            
            # Show total changes prominently
            total_changes = cmp.get("total_changes", 0)
            if total_changes > 0:
                st.warning(f"Найдено {total_changes} изменений в документе")
            else:
                st.success("Документы идентичны или практически идентичны")
            
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Схожесть", f"{cmp['similarity']}%")
            c2.metric("Добавлено", cmp["added_count"])
            c3.metric("Удалено", cmp["removed_count"])
            c4.metric("Изменено", cmp.get("modified_count", 0))

            # Section-level changes (NEW)
            if cmp.get("section_changes"):
                st.subheader("Изменения по разделам (статьи, главы)")
                sec = cmp["section_changes"]
                
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("Всего разделов (старая)", sec.get("total_sections_old", 0))
                c2.metric("Всего разделов (новая)", sec.get("total_sections_new", 0))
                c3.metric("Добавлено разделов", len(sec.get("added", [])))
                c4.metric("Удалено разделов", len(sec.get("removed", [])))
                
                if sec.get("added"):
                    with st.expander(f"Добавленные разделы ({len(sec['added'])})", expanded=True):
                        for s in sec['added'][:10]:
                            st.success(f"**{s['number']}** - {s['title']}")
                            if s.get('content'):
                                st.caption(s['content'][:150] + "...")
                
                if sec.get("removed"):
                    with st.expander(f"Удаленные разделы ({len(sec['removed'])})", expanded=True):
                        for s in sec['removed'][:10]:
                            st.error(f"**{s['number']}** - {s['title']}")
                            if s.get('content'):
                                st.caption(s['content'][:150] + "...")
                
                if sec.get("modified"):
                    with st.expander(f"Измененные разделы ({len(sec['modified'])})", expanded=True):
                        for s in sec['modified'][:10]:
                            st.warning(f"**{s['number']}** - {s['title']} (схожесть: {s['similarity']}%)")
                            col1, col2 = st.columns(2)
                            with col1:
                                st.caption("Было:")
                                st.text(s.get('old_content', '')[:150] + "...")
                            with col2:
                                st.caption("Стало:")
                                st.text(s.get('new_content', '')[:150] + "...")
                st.markdown("---")

            compare_df = pd.DataFrame({
                "Тип": ["Добавлено", "Удалено", "Изменено"],
                "Количество": [cmp["added_count"], cmp["removed_count"], cmp.get("modified_count", 0)]
            })
            fig_cmp = px.bar(compare_df, x="Тип", y="Количество", title="Распределение изменений (строки)", 
                            color="Тип", color_discrete_map={
                                "Добавлено": "#10b981",
                                "Удалено": "#ef4444", 
                                "Изменено": "#f59e0b"
                            })
            st.plotly_chart(fig_cmp, use_container_width=True)

            if cmp.get("modified_examples"):
                st.write(f"**Измененные фрагменты** (показано {len(cmp['modified_examples'])} из {cmp.get('modified_count', 0)})")
                mod_rows = []
                for m in cmp["modified_examples"][:20]:  # Limit to 20 for performance
                    mod_rows.append({
                        "До": m["before"][:300],
                        "После": m["after"][:300],
                        "Схожесть %": m["similarity"]
                    })
                st.dataframe(pd.DataFrame(mod_rows), use_container_width=True, hide_index=True)

            if cmp["added_examples"]:
                with st.expander(f"Добавленные фрагменты ({len(cmp['added_examples'])} из {cmp['added_count']})", expanded=False):
                    st.code("\n".join(cmp["added_examples"]))  # Show all
            if cmp["removed_examples"]:
                with st.expander(f"Удаленные фрагменты ({len(cmp['removed_examples'])} из {cmp['removed_count']})", expanded=False):
                    st.code("\n".join(cmp["removed_examples"]))  # Show all
            if cmp.get("diff_preview"):
                with st.expander("Линейный diff (фрагмент)", expanded=False):
                    st.code("\n".join(cmp["diff_preview"]))  # Show all
            if cmp.get("visual_blocks"):
                st.write("**Визуальное сравнение построчно**")
                st.caption("Зелёный = добавлено | Красный = удалено | Жёлтый = изменено | Серый = без изменений")
                _render_visual_diff_blocks(cmp["visual_blocks"])


def show_demo_examples(use_ai: bool, ai_provider: str = None):
    """Show demo examples."""
    st.header("Демо примеры")
    
    examples = {
        "Пример 1: Устаревшее министерство": {
            "title": "О Министерстве связи и информации",
            "text": """
ЗАКОН РЕСПУБЛИКИ КАЗАХСТАН

Статья 1. Министерство связи и информации Республики Казахстан является 
центральным исполнительным органом, осуществляющим руководство в сфере связи.

Статья 2. Министерство осуществляет:
- разработку государственной политики в области связи
- регулирование деятельности операторов связи

Примечание: Министерство связи и информации было упразднено Указом Президента 
от 6 августа 2019 года, функции переданы Министерству цифрового развития.
            """
        },
        "Пример 2: Противоречие": {
            "title": "О государственном контроле",
            "text": """
Статья 5. Проверки проводятся не чаще одного раза в год.

Статья 12. Внеплановые проверки могут проводиться в любое время 
без ограничений по количеству.

Статья 18. Повторные проверки проводятся не реже двух раз в год.
            """
        }
    }
    
    selected = st.selectbox("Выберите пример", list(examples.keys()))
    
    example = examples[selected]
    
    st.subheader(example['title'])
    st.text_area("Текст", example['text'], height=200, disabled=True)
    
    if st.button("Анализировать пример", type="primary"):
        document = {
            "id": "demo",
            "title": example['title'],
            "full_text": example['text'],
            "doc_type": "text",
            "source": "demo"
        }
        run_analysis_pipeline(document, use_ai, ai_provider)


def show_search_interface(use_ai: bool, ai_provider: str = None):
    """Show search interface (stable adilet-only behavior)."""
    st.header("Поиск в базе adilet.zan.kz")

    if "search_docs" not in st.session_state:
        st.session_state.search_docs = []
    
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        query = st.text_input("Поисковый запрос", placeholder="министерство")
    
    with col2:
        year_from = st.number_input("Год от", min_value=1991, max_value=2035, value=2010)
    
    with col3:
        year_to = st.number_input("Год до", min_value=1991, max_value=2035, value=datetime.now().year)

    with col4:
        limit = st.number_input("Количество", min_value=5, max_value=100, value=30, step=5)

    with col5:
        max_pages = st.number_input("Страниц", min_value=1, max_value=50, value=10, step=1)
    use_local_relevance = st.checkbox(
        "Включить локальный фильтр релевантности (экспериментально)",
        value=False,
        help="Выключено по умолчанию: показывается нативная выдача Adilet как в старой версии."
    )
    hide_procedural_acts = st.checkbox(
        "Скрывать процессуальные акты (о рассмотрении на соответствие Конституции)",
        value=False,
        help="Включите, чтобы убрать служебные акты Конституционного Суда из выдачи"
    )
    
    if st.button("Найти документы"):
        scraper = AdiletScraper()
        st.session_state.search_docs = scraper.search_documents(
            query=query,
            year_from=year_from,
            year_to=year_to,
            limit=limit,
            max_pages=max_pages,
        )
    
    docs = st.session_state.search_docs
    should_postprocess = use_local_relevance or hide_procedural_acts

    if should_postprocess:
        filtered_docs = _postprocess_search_results(
            docs,
            query=query,
            hide_procedural_acts=hide_procedural_acts
        )
        # Safety fallback: never hide all results when source has documents.
        if docs and not filtered_docs:
            filtered_docs = docs
            st.warning("Локальный фильтр дал пустую выборку — показана исходная выдача adilet.zan.kz.")
    else:
        filtered_docs = docs

    if filtered_docs:
        hidden_count = max(0, len(docs) - len(filtered_docs))
        st.success(f"Найдено {len(filtered_docs)} документов")
        if hidden_count > 0:
            st.caption(f"Скрыто результатов после локальной фильтрации: {hidden_count}")
        if should_postprocess:
            st.caption("⚠ Применена локальная пост-обработка результатов")
        else:
            st.caption("✓ Показана нативная выдача поиска adilet.zan.kz (как в предыдущих версиях)")
        render_search_analytics(filtered_docs)

        st.markdown("---")
        st.subheader("Список найденных документов")
        per_page = st.selectbox("Документов на странице", [10, 20, 30, 50], index=1)
        total_docs = len(filtered_docs)
        total_pages = max(1, (total_docs + per_page - 1) // per_page)
        page = st.number_input("Страница", min_value=1, max_value=total_pages, value=1, step=1)
        start_idx = (page - 1) * per_page
        end_idx = min(start_idx + per_page, total_docs)
        page_docs = filtered_docs[start_idx:end_idx]
        st.caption(f"Показано {len(page_docs)} из {total_docs} документов (страница {page}/{total_pages})")

        scraper = AdiletScraper()
        for i, doc in enumerate(page_docs, start=start_idx):
            title = doc.get('title', 'Без названия')
            date = doc.get('date', 'Дата неизвестна')
            doc_key = doc.get('id', f"doc_{i}")

            with st.expander(f"{title} ({date})"):
                if 'doc_number' in doc:
                    st.write(f"**Номер:** {doc['doc_number']}")
                if 'type' in doc:
                    st.write(f"**Тип:** {doc['type']}")
                if 'status' in doc:
                    st.write(f"**Статус:** {doc['status']}")
                st.write(f"**URL:** {doc.get('url', 'N/A')}")
                
                if st.button("Анализировать", key=f"analyze_{doc_key}_{i}"):
                    with st.spinner("Загрузка полного текста..."):
                        full_doc = scraper.fetch_document(doc['url'])
                        if not full_doc:
                            st.session_state.last_status = "Ошибка загрузки документа"
                            st.error("Не удалось загрузить полный текст документа с adilet.zan.kz")
                            continue
                        doc.update(full_doc)
                        doc["doc_type"] = "html"
                        doc["source"] = "adilet"
                        run_analysis_pipeline(doc, use_ai, ai_provider)
    else:
        st.info("Введите запрос и нажмите 'Найти документы'")


def show_multi_analysis(use_ai: bool, ai_provider: str = None):
    """Analyze multiple uploaded documents in one run."""
    st.header("Мульти-анализ документов")
    files = st.file_uploader(
        "Загрузите несколько документов",
        type=["txt", "pdf", "docx"],
        accept_multiple_files=True,
        help="Можно загрузить сразу несколько файлов для пакетного анализа.",
    )
    if not files:
        st.info("Загрузите 2+ файла для пакетного анализа.")
        return

    st.caption(f"Загружено файлов: {len(files)}")
    if st.button("Запустить мульти-анализ", type="primary"):
        summary_rows = []
        detail_results = []
        total_issues = 0
        high_risk_docs = 0

        for idx, f in enumerate(files, start=1):
            with st.spinner(f"Анализ {idx}/{len(files)}: {f.name}"):
                text = extract_text_from_file(f)
                if not text.strip():
                    summary_rows.append(
                        {
                            "Документ": f.name,
                            "Проблемы": 0,
                            "Критичность": 0,
                            "Риск": 0,
                            "Статус": "Ошибка чтения",
                        }
                    )
                    continue

                document = {
                    "id": f"batch_{idx}",
                    "title": f.name,
                    "full_text": text,
                    "doc_type": Path(f.name).suffix.lstrip(".").lower(),
                    "source": "batch_upload",
                }
                ai_client = None
                if use_ai and ai_provider and ai_provider != "none":
                    runtime_key = st.session_state.get("runtime_ai_key", "")
                    ai_client = AIClient(provider=ai_provider, api_key=runtime_key if runtime_key else None)
                detector = DeadRegulationDetector(ai_client=ai_client)
                result = detector.analyze_document(document)

                issues_count = len(result.get("issues_found", []))
                severity = result.get("severity_score", 0)
                risk = result.get("risk_assessment", {}).get("overall_risk", 0)
                total_issues += issues_count
                if risk >= 50:
                    high_risk_docs += 1

                summary_rows.append(
                    {
                        "Документ": f.name,
                        "Проблемы": issues_count,
                        "Критичность": severity,
                        "Риск": risk,
                        "Статус": "Готово",
                    }
                )
                detail_results.append(
                    {
                        "document_name": f.name,
                        "result": result,
                    }
                )

                get_stats_service().track_document(
                    _build_stats_event(
                        document_type=document["doc_type"],
                        source=document["source"],
                        ai_result="success" if result.get("ai_used", False) else "skipped",
                        document_title=document["title"],
                        issues_found=issues_count,
                        analysis_status="success",
                    )
                )

        st.success("Мульти-анализ завершен")
        c1, c2, c3 = st.columns(3)
        c1.metric("Документов", len(files))
        c2.metric("Всего проблем", total_issues)
        c3.metric("Высокий риск (>=50)", high_risk_docs)
        st.dataframe(pd.DataFrame(summary_rows), use_container_width=True, hide_index=True)

        st.markdown("---")
        st.subheader("Детализация по документам")
        for idx, item in enumerate(detail_results, start=1):
            doc_name = item["document_name"]
            result = item["result"]
            issues = result.get("issues_found", [])
            risk = result.get("risk_assessment", {})
            deps = result.get("law_dependencies", [])

            with st.expander(f"{idx}. {doc_name} — проблем: {len(issues)}", expanded=False):
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("Проблем", len(issues))
                m2.metric("Критичность", f"{result.get('severity_score', 0)}/100")
                m3.metric("Риск", f"{risk.get('overall_risk', 0)}/100")
                m4.metric("Зависимости", len(deps))

                if issues:
                    max_issues_to_show = min(25, len(issues))
                    st.caption(f"Показано {max_issues_to_show} из {len(issues)} проблем")
                    for j, issue in enumerate(issues[:max_issues_to_show], start=1):
                        st.markdown(f"**Проблема {j}: {issue_type_ru(issue.get('type', 'unknown'))}**")
                        st.write(f"**Критичность:** {severity_ru(issue.get('severity', 'Low'))}")
                        st.write(f"**Объяснение:** {issue.get('explanation', 'Не указано')}")
                        st.write(f"**Рекомендация:** {issue.get('recommendation', 'Не указано')}")
                        quote_text = issue.get("quote", "")
                        if quote_text:
                            st.text_area(
                                f"Цитата {j}",
                                value=quote_text[:1200],
                                height=120,
                                disabled=True,
                                key=f"multi_quote_{idx}_{j}_{abs(hash(doc_name)) % 100000}",
                            )
                        st.markdown("---")
                else:
                    st.success("Проблем не обнаружено")


def show_system_analytics():
    """Render global analytics dashboard from persisted stats."""
    st.header("Аналитика системы")
    stats = get_stats_service().get_stats()

    total_docs = int(stats.get("total_processed_documents", 0))
    ai_results = stats.get("ai_results", {})
    analysis_results = stats.get("analysis_results", {})
    sources = stats.get("document_sources", {})
    doc_types = stats.get("document_types", {})
    recent = stats.get("recent_analyses", [])

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Всего анализов", total_docs)
    c2.metric("Успешных AI-вызовов", int(ai_results.get("success", 0)))
    c3.metric("Ошибки анализа", int(analysis_results.get("failed", 0)))
    c4.metric("Источников", len(sources))

    success = int(analysis_results.get("success", 0))
    failed = int(analysis_results.get("failed", 0))
    total_runs = max(1, success + failed)
    ai_success = int(ai_results.get("success", 0))
    ai_failed = int(ai_results.get("failed", 0))
    ai_skipped = int(ai_results.get("skipped", 0))
    ai_total = max(1, ai_success + ai_failed + ai_skipped)

    q1, q2, q3, q4 = st.columns(4)
    q1.metric("Успешность анализа", f"{round(success * 100 / total_runs, 1)}%")
    q2.metric("AI ошибок", ai_failed)
    q3.metric("AI пропущено", ai_skipped)
    q4.metric("Доля AI-success", f"{round(ai_success * 100 / ai_total, 1)}%")

    if sources:
        src_df = pd.DataFrame({"Источник": list(sources.keys()), "Количество": list(sources.values())})
        fig_src = px.pie(src_df, names="Источник", values="Количество", title="Распределение по источникам")
        st.plotly_chart(fig_src, use_container_width=True)

    if doc_types:
        types_df = pd.DataFrame({"Тип": list(doc_types.keys()), "Количество": list(doc_types.values())})
        fig_types = px.bar(types_df, x="Тип", y="Количество", title="Распределение по типам документов")
        st.plotly_chart(fig_types, use_container_width=True)

    st.subheader("Последние анализы")
    if recent:
        recent_df = pd.DataFrame(list(reversed(recent[-100:])))
        st.dataframe(recent_df, use_container_width=True, hide_index=True)
        if "issues_found" in recent_df.columns:
            st.subheader("Динамика найденных проблем (последние анализы)")
            trend_df = recent_df.copy().head(30).iloc[::-1]
            trend_df["Индекс"] = list(range(1, len(trend_df) + 1))
            fig_trend = px.line(
                trend_df,
                x="Индекс",
                y="issues_found",
                markers=True,
                title="Проблемы по последним анализам",
                labels={"issues_found": "Найдено проблем"},
            )
            st.plotly_chart(fig_trend, use_container_width=True)
    else:
        st.info("История анализов пока пуста.")


def _postprocess_search_results(docs: list, query: str = "", hide_procedural_acts: bool = True) -> list:
    """Remove noisy procedural acts and exact duplicates from search output."""
    if not docs:
        return []

    cleaned = []
    seen_titles = set()
    query_norm = re.sub(r"\s+", " ", (query or "").lower()).strip()
    query_tokens = [t for t in re.findall(r"[а-яёa-z]{3,}", query_norm) if t not in {
        "о", "об", "и", "или", "в", "во", "на", "по", "для", "с", "со", "к", "ко",
        "от", "до", "за", "из", "у", "а", "не", "при", "под", "над", "без"
    }]
    
    constitutional_intent = any(
        token in query_norm for token in ("конституц", "конституцион", "конституционный суд")
    )
    
    noise_patterns = [
        r"^о рассмотрении на соответствие конституции",
        r"^о проверке конституционности",
        r"^об официальном толковании",
        r"^нормативное постановление конституционного суда .*о рассмотрении на соответствие конституции",
        r"^о состоянии конституционной законности",
        r"^послание конституционного совета",
        r"^об утверждении регламента конституционного суда",
        r"^о пересмотре некоторых нормативных постановлений конституционного совета",
    ]

    # Score and filter documents
    scored_docs = []
    for doc in docs:
        title = str(doc.get("title", "")).strip()
        if not title:
            continue

        normalized_title = re.sub(r"\s+", " ", title.lower()).strip()
        
        # Skip exact duplicates
        if normalized_title in seen_titles:
            continue
        seen_titles.add(normalized_title)

        # Skip procedural acts (unless searching for constitutional topics)
        if hide_procedural_acts and not constitutional_intent:
            if any(re.search(pattern, normalized_title) for pattern in noise_patterns):
                continue

        # Calculate relevance score
        score = 0
        
        # Exact phrase match = highest score
        if query_norm and query_norm in normalized_title:
            score += 1000
        
        # Match query tokens
        matched_tokens = 0
        for token in query_tokens:
            # Exact token match
            if token in normalized_title:
                matched_tokens += 1
                score += 100
            # Prefix match (for word forms)
            elif len(token) >= 5 and any(token[:5] in word for word in normalized_title.split()):
                matched_tokens += 1
                score += 50
        
        # All tokens matched = bonus
        if query_tokens and matched_tokens == len(query_tokens):
            score += 200
        
        # Boost primary document types
        if re.search(r"\b(кодекс|закон|конституция)\b", normalized_title):
            # But only if it's not buried in procedural language
            if not re.search(r"^о (рассмотрении|проверке|внесении изменений)", normalized_title):
                score += 150
        
        # Penalize secondary/procedural documents
        if re.search(r"^о (рассмотрении|проверке|внесении изменений|состоянии)", normalized_title):
            score -= 300
        
        # Skip documents with zero relevance (unless no query)
        if query_norm and score <= 0:
            continue
        
        doc_copy = dict(doc)
        doc_copy['_relevance_score'] = score
        scored_docs.append(doc_copy)

    # Sort by relevance score (highest first)
    scored_docs.sort(key=lambda d: d.get('_relevance_score', 0), reverse=True)
    
    # Remove score from output
    for doc in scored_docs:
        doc.pop('_relevance_score', None)
    
    return scored_docs


def build_dependency_graph(doc_title: str, dependencies: List[dict]):
    """Build interactive dependency graph using Plotly."""
    import math
    
    # Create graph data structures
    nodes = [{'id': 'doc', 'label': doc_title[:50], 'type': 'document', 'size': 30, 'color': 'red'}]
    edges = []
    
    # Add dependency nodes
    for i, dep in enumerate(dependencies[:30]):  # Limit to 30 for readability
        node_id = f"dep_{i}"
        label = dep['reference'][:40]
        dep_type = dep['type']
        
        # Size and color by type
        if dep_type in ['law_reference', 'code_reference']:
            size = 20
            color = 'blue' if dep_type == 'code_reference' else 'teal'
        else:
            size = 12
            color = 'orange' if dep_type == 'article_reference' else 'green'
        
        nodes.append({'id': node_id, 'label': label, 'type': dep_type, 'size': size, 'color': color})
        edges.append({'source': 'doc', 'target': node_id})
    
    # Create Plotly figure
    fig = go.Figure()
    
    # Calculate positions (circular layout)
    n = len(nodes)
    for i, node in enumerate(nodes):
        if node['id'] == 'doc':
            x, y = 0, 0  # Center
        else:
            angle = 2 * math.pi * (i - 1) / (n - 1)
            x = math.cos(angle)
            y = math.sin(angle)
        node['x'] = x
        node['y'] = y
    
    # Add edges
    edge_trace = []
    for edge in edges:
        source = next(n for n in nodes if n['id'] == edge['source'])
        target = next(n for n in nodes if n['id'] == edge['target'])
        edge_trace.append(go.Scatter(
            x=[source['x'], target['x'], None],
            y=[source['y'], target['y'], None],
            mode='lines',
            line=dict(width=0.5, color='#888'),
            hoverinfo='none',
            showlegend=False
        ))
    
    for trace in edge_trace:
        fig.add_trace(trace)
    
    # Add nodes
    for node in nodes:
        fig.add_trace(go.Scatter(
            x=[node['x']],
            y=[node['y']],
            mode='markers+text',
            marker=dict(size=node['size'], color=node['color']),
            text=node['label'],
            textposition='top center',
            hovertext=f"{node['type']}: {node['label']}",
            hoverinfo='text',
            showlegend=False
        ))
    
    fig.update_layout(
        title='Граф правовых зависимостей',
        showlegend=False,
        hovermode='closest',
        margin=dict(b=0, l=0, r=0, t=40),
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        height=600
    )
    
    return fig


def display_analysis_results(result: dict):
    """Display analysis results."""
    st.markdown("---")
    st.header("Результаты анализа")

    # AI transparency block
    if "ai_used" in result:
        ai_used = result.get("ai_used", False)
        if ai_used:
            provider = result.get("ai_provider", "unknown")
            model = result.get("ai_model") or "unknown"
            latency = result.get("ai_latency_sec", 0.0)
            ai_added = result.get("ai_issues_added", 0)
            st.success(f"AI анализ выполнен: {provider} ({model})")
            c1, c2, c3 = st.columns(3)
            with c1:
                st.metric("AI вызов", "Да")
            with c2:
                st.metric("Время AI", f"{latency} сек")
            with c3:
                st.metric("Проблем от AI", ai_added)
        else:
            st.warning("AI не использовался в этом анализе (работали только rule-based проверки)")
    else:
        st.warning("AI блок не вернул статус. Проверьте ключ/API провайдера.")
    
    # Metrics
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Найдено проблем", len(result['issues_found']))
    
    with col2:
        st.metric("Оценка критичности", f"{result['severity_score']}/100")
    
    with col3:
        critical = "Да" if result['has_critical_issues'] else "Нет"
        st.metric("Критические проблемы", critical)
    
    # Issues
    if result['issues_found']:
        st.subheader("Выявленные проблемы")

        severity_counts = Counter(severity_ru(issue.get("severity", "Low")) for issue in result["issues_found"])
        type_counts = Counter(issue.get("type", "unknown") for issue in result["issues_found"])

        chart_col1, chart_col2 = st.columns(2)
        with chart_col1:
            sev_df = pd.DataFrame({
                "severity": list(severity_counts.keys()),
                "count": list(severity_counts.values())
            })
            fig_sev = px.bar(sev_df, x="severity", y="count", title="Распределение по критичности")
            st.plotly_chart(fig_sev, use_container_width=True)

        with chart_col2:
            type_df = pd.DataFrame({
                "type": list(type_counts.keys()),
                "count": list(type_counts.values())
            })
            fig_type = px.pie(type_df, names="type", values="count", title="Типы найденных проблем")
            st.plotly_chart(fig_type, use_container_width=True)
        
        for i, issue in enumerate(result['issues_found'], 1):
            severity_en = issue.get('severity', 'Low')
            severity = severity_ru(severity_en)
            severity_class = f"{severity_en.lower()}-severity"
            
            with st.container():
                st.markdown(f"""
                <div class="issue-card {severity_class}">
                    <h4>Проблема {i}: {issue_type_ru(issue.get('type', 'unknown'))}</h4>
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    st.write("**Цитата:**")
                    quote_text = issue.get('quote', 'N/A')
                    quote_lines = max(10, min(40, len(str(quote_text)) // 90))
                    st.text_area(
                        label="Текст цитаты",
                        value=quote_text,
                        height=quote_lines * 24,
                        disabled=True,
                        key=f"quote_{i}_{hash(str(quote_text)) % 100000}"
                    )
                    
                    st.write("**Объяснение:**")
                    st.write(issue.get('explanation', 'N/A'))
                
                with col2:
                    st.write(f"**Критичность:** {severity}")
                    st.write("**Рекомендация:**")
                    st.info(issue.get('recommendation', 'Требуется ревизия нормы'))
                    refs = build_reference_links(issue)
                    if refs:
                        st.write("**Связанные нормы:**")
                        for ref in refs:
                            st.markdown(f"- [{ref['label']}]({ref['url']})")
                
                st.markdown("---")
    
    else:
        st.success("Критических проблем не обнаружено")
    
    # AI Summary
    if 'ai_note' in result:
        st.info(result['ai_note'])
    if 'ai_error' in result:
        st.warning(_format_ai_error_message(result['ai_error'], result.get("ai_summary", "")))
    if 'ai_summary' in result:
        st.subheader("AI резюме")
        
        # Show chunks info if available
        if 'ai_metadata' in result:
            meta = result['ai_metadata']
            if meta.get('chunks_analyzed', 0) > 1:
                st.info(f"Документ проанализирован по частям: {meta['chunks_analyzed']} частей, всего символов: {meta.get('total_length', 'N/A'):,}")
        
        st.write(result['ai_summary'])
    
    # Risk Assessment Section
    risk = result.get('risk_assessment', {}) or {}
    categories = risk.get('categories', {})

    st.markdown("---")
    st.subheader("Оценка правовых рисков")
    rc1, rc2, rc3 = st.columns(3)
    with rc1:
        st.metric("Общий риск", f"{risk.get('overall_risk', 0)}/100")
    with rc2:
        st.metric("Уровень риска", risk.get('risk_label', 'Неизвестно'))
    with rc3:
        st.metric("Категории", len(categories))

    if categories:
        cat_data = {
            'Категория': [cat.get('label', '') for cat in categories.values()],
            'Балл': [cat.get('score', 0) for cat in categories.values()]
        }
        fig_risk = px.bar(
            pd.DataFrame(cat_data),
            x='Категория',
            y='Балл',
            title='Риски по категориям',
            color='Балл',
            color_continuous_scale=['green', 'yellow', 'orange', 'red']
        )
        st.plotly_chart(fig_risk, use_container_width=True)

    recommendations = risk.get('recommendations', [])
    if recommendations:
        st.write("**Рекомендации:**")
        for rec in recommendations:
            priority = rec.get('priority', 'Medium')
            st.write(f"**{priority_ru(priority)}**: {rec.get('text', '')}")
    else:
        st.info("Рекомендации пока не требуются для этого документа.")

    # Law Dependencies Section
    dependencies = result.get('law_dependencies', []) or []
    dep_types = Counter(d.get('type', 'unknown') for d in dependencies)
    type_labels = {
        'law_reference': 'Законы',
        'code_reference': 'Кодексы',
        'article_reference': 'Статьи',
        'section_reference': 'Пункты'
    }

    st.markdown("---")
    st.subheader("Правовые зависимости")
    dc1, dc2, dc3 = st.columns(3)
    with dc1:
        st.metric("Всего ссылок", len(dependencies))
    with dc2:
        st.metric("Типов ссылок", len(dep_types))
    with dc3:
        law_refs = sum(1 for d in dependencies if d.get('type') in ['law_reference', 'code_reference'])
        st.metric("Ссылки на законы", law_refs)

    if dependencies:
        dep_type_data = {
            'Тип': [type_labels.get(t, t) for t in dep_types.keys()],
            'Количество': list(dep_types.values())
        }
        fig_deps = px.pie(
            pd.DataFrame(dep_type_data),
            names='Тип',
            values='Количество',
            title='Распределение ссылок по типам'
        )
        st.plotly_chart(fig_deps, use_container_width=True)

        st.write("**Граф зависимостей:**")
        graph_fig = build_dependency_graph(result.get('document_title', 'Документ'), dependencies)
        st.plotly_chart(graph_fig, use_container_width=True)

        st.write("**Список зависимостей:**")
        dep_df = pd.DataFrame([{
            'Тип': type_labels.get(d.get('type', ''), d.get('type', '')),
            'Ссылка': str(d.get('reference', ''))[:80],
            'Контекст': (str(d.get('context', ''))[:100] + '...') if d.get('context') else ''
        } for d in dependencies[:20]])
        st.dataframe(dep_df, use_container_width=True)
    else:
        st.info("В документе пока не найдено явных правовых ссылок.")

    # Timeline Section
    timeline = result.get('timeline_points', []) or []
    st.markdown("---")
    st.subheader("Временная шкала документа")

    tc1, tc2, tc3 = st.columns(3)
    with tc1:
        st.metric("Упоминаний годов", len(timeline))
    with tc2:
        years = [t.get('year') for t in timeline if t.get('year') is not None]
        st.metric("Период", f"{min(years)} - {max(years)}" if years else "N/A")
    with tc3:
        recent_count = sum(1 for t in timeline if int(t.get('year', 0)) >= 2020)
        st.metric("С 2020 года", recent_count)

    if timeline:
        year_counts = Counter(t.get('year') for t in timeline if t.get('year') is not None)
        timeline_data = {
            'Год': list(year_counts.keys()),
            'Упоминаний': list(year_counts.values())
        }
        fig_timeline = px.line(
            pd.DataFrame(timeline_data),
            x='Год',
            y='Упоминаний',
            title='Упоминания по годам',
            markers=True
        )
        st.plotly_chart(fig_timeline, use_container_width=True)
    else:
        st.info("Не удалось извлечь годы из текста документа.")

    render_download_reports(result)


def init_session_state() -> None:
    """Initialize persistent UI counters and status."""
    defaults = {
        "analyzed_docs": 0,
        "total_issues": 0,
        "last_status": "Ожидание",
        "last_result": None
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def run_analysis_pipeline(document: dict, use_ai: bool, ai_provider: str = None) -> None:
    """Run analysis with explicit step-by-step feedback."""
    status_box = st.empty()
    progress = st.progress(0)
    started = time.time()

    try:
        status_box.info("Шаг 1/4: Подготовка документа")
        progress.progress(20)

        ai_client = None
        ai_mode = ai_provider if use_ai else "none"
        if use_ai and ai_provider and ai_provider != "none":
            status_box.info("Шаг 2/4: Инициализация AI клиента")
            runtime_key = st.session_state.get("runtime_ai_key", "")
            ai_client = AIClient(provider=ai_provider, api_key=runtime_key if runtime_key else None)
            if ai_client.is_available():
                st.caption(f"AI активен: {ai_provider} / {getattr(ai_client, 'model', 'unknown')}")
            else:
                st.caption(f"AI не активен: {ai_provider}")
        else:
            status_box.info("Шаг 2/4: AI отключен (none)")

        status_box.info("Шаг 3/4: Анализ текста")
        progress.progress(60)
        detector = DeadRegulationDetector(ai_client=ai_client)
        result = detector.analyze_document(document)

        status_box.info("Шаг 4/4: Формирование отчета")
        progress.progress(90)
        
        st.session_state.analyzed_docs += 1
        st.session_state.total_issues += len(result.get("issues_found", []))
        st.session_state.last_result = result
        st.session_state.last_document_title = document.get("title", "Документ")
        elapsed = round(time.time() - started, 2)
        st.session_state.last_status = f"Успешно ({elapsed} c)"
        
        # Clear progress indicators before showing results
        progress.progress(100)
        status_box.success(f"✓ Анализ завершен за {elapsed} сек. Найдено проблем: {len(result.get('issues_found', []))}")
        
        # Show results AFTER clearing progress
        display_analysis_results(result)

        ai_result = "skipped"
        if ai_mode != "none":
            ai_result = "success" if result.get("ai_used", False) and "ai_error" not in result else "failed"
        get_stats_service().track_document(
            _build_stats_event(
                document_type=document.get("doc_type", "unknown"),
                source=document.get("source", "unknown"),
                ai_result=ai_result,
                document_title=document.get("title", "Без названия"),
                issues_found=len(result.get("issues_found", [])),
                analysis_status="success",
            )
        )
    except Exception as e:
        get_stats_service().track_document(
            _build_stats_event(
                document_type=document.get("doc_type", "unknown"),
                source=document.get("source", "unknown"),
                ai_result="failed" if use_ai else "skipped",
                document_title=document.get("title", "Без названия"),
                issues_found=0,
                analysis_status="failed",
                error_message=str(e),
            )
        )
        st.session_state.last_status = f"Ошибка: {e}"
        status_box.error(f"Анализ не завершен: {e}")


def _build_stats_event(**kwargs) -> StatsEvent:
    """Build StatsEvent with backward compatibility across signatures."""
    accepted = set(inspect.signature(StatsEvent).parameters.keys())
    filtered = {k: v for k, v in kwargs.items() if k in accepted}
    return StatsEvent(**filtered)


def extract_text_from_file(uploaded_file) -> str:
    """Extract text using robust extraction module."""
    try:
        result = extract_text_from_uploaded_file(uploaded_file)
        if result.warning:
            st.warning(result.warning)
        return result.text
    except Exception as e:
        st.error(f"Ошибка чтения файла: {e}")
        return ""


def render_search_analytics(docs: list) -> None:
    """Render timeline and duplicate candidates for search results."""
    if not docs:
        return

    years = []
    for d in docs:
        y = extract_year(d.get("date", ""))
        if y:
            years.append(y)
    if years:
        df = pd.DataFrame({"year": years})
        timeline = df.value_counts("year").reset_index(name="count").sort_values("year")
        fig = px.line(timeline, x="year", y="count", markers=True, title="Таймлайн найденных актов")
        st.plotly_chart(fig, use_container_width=True)

    duplicates = detect_similar_documents(docs)
    if duplicates:
        st.subheader("Похожие документы (кандидаты на дублирование)")
        dup_df = pd.DataFrame(duplicates)
        st.dataframe(dup_df, use_container_width=True, hide_index=True)


def _render_visual_diff_blocks(blocks: list) -> None:
    """Render side-by-side visual diff blocks - shows ALL changes."""
    st.markdown(
        """
        <style>
        .diff-container {max-height: 600px; overflow-y: auto; border: 1px solid #e5e7eb; border-radius: 8px; padding: 12px;}
        .diff-row {display:grid;grid-template-columns:1fr 1fr;gap:10px;margin:6px 0;}
        .diff-cell {
            padding:12px;
            border-radius:6px;
            border:1px solid #cbd5e1;
            white-space:pre-wrap;
            font-family: 'Courier New', monospace;
            font-size: 13px;
            line-height: 1.6;
        }
        .diff-equal {
            background:#f1f5f9;
            border-color:#cbd5e1;
            color:#334155;
        }
        .diff-added {
            background:#dcfce7;
            border-color:#22c55e;
            color:#14532d;
            font-weight: 600;
        }
        .diff-removed {
            background:#fee2e2;
            border-color:#ef4444;
            color:#7f1d1d;
            font-weight: 600;
        }
        .diff-modified {
            background:#fef3c7;
            border-color:#f59e0b;
            color:#78350f;
            font-weight: 600;
        }
        .diff-cell:empty::before {
            content: "—";
            color: #94a3b8;
            font-style: italic;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    
    # Show count info
    st.caption(f"Показано {len(blocks)} строк изменений")
    
    # Add headers
    st.markdown(
        """
        <div class="diff-container">
        <div class="diff-row" style="font-weight:bold;margin-bottom:12px;position:sticky;top:0;background:white;padding:8px 0;z-index:10;">
          <div style="text-align:center;color:#1e293b;">Старая версия</div>
          <div style="text-align:center;color:#1e293b;">Новая версия</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    # Show ALL blocks - no limit!
    for block in blocks:
        btype = block.get("type", "equal")
        if btype == "modified":
            left = block.get("old_html") or block.get("old", "")
            right = block.get("new_html") or block.get("new", "")
        else:
            left = block.get("old", "")
            right = block.get("new", "")
            left = left.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            right = right.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        
        st.markdown(
            f"""
            <div class="diff-row">
              <div class="diff-cell diff-{btype}">{left}</div>
              <div class="diff-cell diff-{btype}">{right}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    
    # Close container
    st.markdown("</div>", unsafe_allow_html=True)


def extract_year(date_text: str):
    m = re.search(r"(19|20)\d{2}", str(date_text))
    return int(m.group(0)) if m else None


def detect_similar_documents(docs: list, threshold: float = 0.82) -> list:
    """Find likely duplicate documents by title similarity."""
    pairs = []
    for i in range(len(docs)):
        for j in range(i + 1, len(docs)):
            t1 = docs[i].get("title", "")
            t2 = docs[j].get("title", "")
            if not t1 or not t2:
                continue
            score = SequenceMatcher(None, t1.lower(), t2.lower()).ratio()
            if score >= threshold:
                pairs.append({
                    "doc_a": t1[:90],
                    "doc_b": t2[:90],
                    "similarity": round(score * 100, 1)
                })
    return sorted(pairs, key=lambda x: x["similarity"], reverse=True)[:20]


def build_reference_links(issue: dict) -> list:
    """Generate links to potentially relevant legal norms."""
    refs = []
    issue_type = issue.get("type", "")
    quote = str(issue.get("quote", "")).lower()
    mapping = {
        "outdated_terms": "https://adilet.zan.kz/rus/search/docs?q=реорганизация+государственных+органов",
        "contradiction": "https://adilet.zan.kz/rus/search/docs?sort_field=dt&sort_desc=true",
        "duplication": "https://adilet.zan.kz/rus/search/docs?q=изменения+и+дополнения",
        "inapplicability": "https://adilet.zan.kz/rus/search/docs?q=порядок+реализации"
    }
    if issue_type in mapping:
        refs.append({"label": f"Поиск по типу: {issue_type}", "url": mapping[issue_type]})
    if "министерство связи и информации" in quote:
        refs.append({
            "label": "Актуальные акты по цифровому развитию",
            "url": "https://adilet.zan.kz/rus/search/docs?q=министерство+цифрового+развития"
        })
    return refs


def render_download_reports(result: dict) -> None:
    """Render download buttons for Excel/PDF/DOCX reports."""
    title = st.session_state.get("last_document_title", "Документ")
    excel_bytes = build_excel_report(result)
    pdf_bytes = build_pdf_report(result, title)
    docx_bytes = build_docx_report(result, title)
    
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button(
            "Скачать отчет Excel",
            data=excel_bytes,
            file_name="dead_regulations_report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    with c2:
        st.download_button(
            "Скачать отчет PDF",
            data=pdf_bytes,
            file_name="dead_regulations_report.pdf",
            mime="application/pdf"
        )
    with c3:
        st.download_button(
            "Скачать отчет DOCX",
            data=docx_bytes,
            file_name="dead_regulations_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


def build_excel_report(result: dict) -> bytes:
    issues = result.get("issues_found", [])
    rows = []
    for idx, issue in enumerate(issues, start=1):
        rows.append({
            "№": idx,
            "Тип": issue.get("type", ""),
            "Критичность": severity_ru(issue.get("severity", "")),
            "Цитата": issue.get("quote", ""),
            "Объяснение": issue.get("explanation", ""),
            "Рекомендация": issue.get("recommendation", "")
        })
    df = pd.DataFrame(rows if rows else [{"№": 1, "Тип": "Нет проблем", "Критичность": "", "Цитата": "", "Объяснение": "", "Рекомендация": ""}])
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Issues")
    output.seek(0)
    return output.read()


def build_pdf_report(result: dict, title: str) -> bytes:
    def _pdf_safe(text: str) -> str:
        return str(text).encode("latin-1", errors="ignore").decode("latin-1")

    try:
        from fpdf import FPDF
    except Exception:
        return b"PDF library unavailable"
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 8, txt=_pdf_safe(f"Dead Regulations Report\nDocument: {title}"))
        pdf.ln(2)
        pdf.multi_cell(0, 8, txt=_pdf_safe(f"Оценка критичности: {result.get('severity_score', 0)}"))
        pdf.multi_cell(0, 8, txt=_pdf_safe(f"Найдено проблем: {len(result.get('issues_found', []))}"))
        pdf.multi_cell(0, 8, txt=_pdf_safe(f"Общий риск: {result.get('risk_assessment', {}).get('overall_risk', 0)}"))
        for i, issue in enumerate(result.get("issues_found", []), start=1):
            pdf.ln(2)
            pdf.set_font("Helvetica", style="B", size=11)
            pdf.multi_cell(0, 8, txt=_pdf_safe(f"{i}. {issue_type_ru(issue.get('type', 'unknown'))} ({severity_ru(issue.get('severity', ''))})"))
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 6, txt=_pdf_safe(f"Цитата: {issue.get('quote', '')[:300]}"))
            pdf.multi_cell(0, 6, txt=_pdf_safe(f"Объяснение: {issue.get('explanation', '')[:300]}"))
            pdf.multi_cell(0, 6, txt=_pdf_safe(f"Рекомендация: {issue.get('recommendation', '')[:300]}"))
        out = pdf.output(dest="S")
        if isinstance(out, bytearray):
            return bytes(out)
        if isinstance(out, str):
            return out.encode("latin-1", errors="ignore")
        return out
    except Exception:
        return b"PDF generation failed"


def build_docx_report(result: dict, title: str) -> bytes:
    """Build professional DOCX report with key analysis sections."""
    try:
        doc = Document()

        heading = doc.add_heading("Отчет по анализу нормативного документа", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Документ: {title}")
        doc.add_paragraph(f"Дата анализа: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.add_heading("Сводные показатели", level=1)
        summary = doc.add_table(rows=4, cols=2)
        summary.style = "Light Grid Accent 1"
        summary.cell(0, 0).text = "Найдено проблем"
        summary.cell(0, 1).text = str(len(result.get("issues_found", [])))
        summary.cell(1, 0).text = "Оценка критичности"
        summary.cell(1, 1).text = f"{result.get('severity_score', 0)}/100"
        summary.cell(2, 0).text = "Критические проблемы"
        summary.cell(2, 1).text = "Да" if result.get("has_critical_issues") else "Нет"
        summary.cell(3, 0).text = "Общий риск"
        summary.cell(3, 1).text = f"{result.get('risk_assessment', {}).get('overall_risk', 0)}/100"

        risk = result.get("risk_assessment", {})
        if risk:
            doc.add_heading("Оценка правовых рисков", level=1)
            doc.add_paragraph(f"Уровень риска: {risk.get('risk_label', 'Неизвестно')}")
            categories = risk.get("categories", {})
            if categories:
                risk_table = doc.add_table(rows=1, cols=2)
                risk_table.style = "Light List Accent 1"
                risk_table.cell(0, 0).text = "Категория"
                risk_table.cell(0, 1).text = "Балл"
                for cat in categories.values():
                    row = risk_table.add_row().cells
                    row[0].text = cat.get("label", "")
                    row[1].text = str(cat.get("score", 0))

            recommendations = risk.get("recommendations", [])
            if recommendations:
                doc.add_paragraph("Рекомендации:")
                for rec in recommendations:
                    doc.add_paragraph(
                        f"{priority_ru(rec.get('priority', 'Medium'))}: {rec.get('text', '')}",
                        style="List Bullet",
                    )

        dependencies = result.get("law_dependencies", [])
        if dependencies:
            doc.add_heading("Правовые зависимости", level=1)
            dep_table = doc.add_table(rows=1, cols=3)
            dep_table.style = "Light Grid Accent 1"
            dep_table.cell(0, 0).text = "Тип"
            dep_table.cell(0, 1).text = "Ссылка"
            dep_table.cell(0, 2).text = "Контекст"
            for dep in dependencies[:30]:
                row = dep_table.add_row().cells
                row[0].text = dep.get("type", "")
                row[1].text = dep.get("reference", "")[:120]
                row[2].text = dep.get("context", "")[:180]

        doc.add_heading("Выявленные проблемы", level=1)
        issues = result.get("issues_found", [])
        if not issues:
            doc.add_paragraph("Проблем не обнаружено.")
        else:
            for idx, issue in enumerate(issues, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{idx}. {issue_type_ru(issue.get('type', 'unknown'))} ({severity_ru(issue.get('severity', ''))})")
                run.bold = True
                if issue.get("severity") == "High":
                    run.font.color.rgb = RGBColor(255, 0, 0)
                elif issue.get("severity") == "Medium":
                    run.font.color.rgb = RGBColor(255, 140, 0)
                doc.add_paragraph(f"Цитата: {issue.get('quote', '')[:400]}")
                doc.add_paragraph(f"Объяснение: {issue.get('explanation', '')[:400]}")
                doc.add_paragraph(f"Рекомендация: {issue.get('recommendation', '')[:300]}")

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    except Exception:
        return b"DOCX generation failed"


if __name__ == "__main__":
    main()
